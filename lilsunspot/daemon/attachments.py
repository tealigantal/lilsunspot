from __future__ import annotations

import base64
import csv
import io
import mimetypes
import re
import secrets
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from . import conversations
from .config_paths import RuntimePaths, ensure_runtime_dirs


ATTACHMENT_DIR_NAME = "attachments"
IMAGE_PREVIEW_MAX_BYTES = 8 * 1024 * 1024
TEXT_READ_MAX_BYTES = 2 * 1024 * 1024
SUMMARY_TOO_LARGE_BYTES = 25 * 1024 * 1024
SUMMARY_MAX_CHARS = 4000
IMAGE_PREVIEW_ONLY_REASON = "图片已收到并可预览；当前还没有完成视觉识别。"

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
TEXT_EXTS = {".txt", ".md", ".log", ".json", ".yaml", ".yml", ".toml"}
CSV_EXTS = {".csv"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx"}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".opus", ".silk"}
VIDEO_EXTS = {".mp4", ".mov", ".avi", ".mkv", ".webm", ".3gp"}
ARCHIVE_EXTS = {".zip", ".rar", ".7z", ".tar", ".gz"}


class AttachmentError(ValueError):
    pass


def _now_yyyymm() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m")


def _safe_name(name: str) -> str:
    base = Path(name).name.replace("\x00", "").strip()
    if not base or base in {".", ".."}:
        base = "attachment"
    base = re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", "_", base)
    base = re.sub(r"\s+", " ", base).strip()
    return base[:160] or "attachment"


def _resolve_file(path: str | Path) -> Path:
    candidate = Path(path).expanduser()
    try:
        resolved = candidate.resolve(strict=True)
    except (OSError, RuntimeError, ValueError) as exc:
        raise AttachmentError("文件已收到，但本地缓存路径不可读取。") from exc
    if not resolved.is_file():
        raise AttachmentError("文件已收到，但本地缓存不是普通文件。")
    return resolved


def _is_within(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False


def _allowed_roots(paths: RuntimePaths) -> list[Path]:
    roots = [paths.data_dir / ATTACHMENT_DIR_NAME]
    for child in (
        paths.hermes_home / "cache",
        paths.hermes_home / "cache" / "images",
        paths.hermes_home / "cache" / "audio",
        paths.hermes_home / "cache" / "videos",
        paths.hermes_home / "cache" / "documents",
        paths.hermes_home / "cache" / "screenshots",
        paths.hermes_home / "image_cache",
        paths.hermes_home / "audio_cache",
        paths.hermes_home / "video_cache",
        paths.hermes_home / "document_cache",
        paths.hermes_home / "browser_screenshots",
    ):
        roots.append(child)
    safe_roots: list[Path] = []
    for root in roots:
        try:
            safe_roots.append(root.expanduser().resolve(strict=False))
        except (OSError, RuntimeError, ValueError):
            continue
    return safe_roots


def assert_safe_attachment_path(path: str | Path, paths: RuntimePaths | None = None) -> Path:
    runtime_paths = paths or ensure_runtime_dirs()
    resolved = _resolve_file(path)
    if any(_is_within(resolved, root) for root in _allowed_roots(runtime_paths)):
        return resolved
    raise AttachmentError("文件已收到，但来源路径不在小黑子的安全缓存目录内。")


def attachment_storage_root(paths: RuntimePaths | None = None) -> Path:
    runtime_paths = paths or ensure_runtime_dirs()
    root = runtime_paths.data_dir / ATTACHMENT_DIR_NAME
    root.mkdir(parents=True, exist_ok=True)
    return root


def is_safe_stored_attachment(path: str | Path, paths: RuntimePaths | None = None) -> Path:
    runtime_paths = paths or ensure_runtime_dirs()
    root = attachment_storage_root(runtime_paths).resolve(strict=False)
    resolved = _resolve_file(path)
    if not _is_within(resolved, root):
        raise AttachmentError("附件只能从小黑子的附件目录打开。")
    return resolved


def _guess_mime(file_path: Path, provided_mime: str = "") -> str:
    provided = provided_mime.strip()
    if provided:
        return provided
    guessed, _ = mimetypes.guess_type(str(file_path))
    return guessed or "application/octet-stream"


def _copy_to_attachment_dir(source: Path, attachment_id: str, paths: RuntimePaths) -> Path:
    target_dir = attachment_storage_root(paths) / _now_yyyymm()
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"attachment_{attachment_id}_{_safe_name(source.name)}"
    target_root = attachment_storage_root(paths).resolve(strict=False)
    resolved_target = target.resolve(strict=False)
    if not _is_within(resolved_target, target_root):
        raise AttachmentError("附件文件名不安全，已拒绝保存。")
    shutil.copy2(source, target)
    return target


def _clip_summary(text: str) -> str:
    text = text.strip()
    if len(text) <= SUMMARY_MAX_CHARS:
        return text
    return text[:SUMMARY_MAX_CHARS].rstrip() + "\n……（摘要已截断）"


def _read_text_prefix(path: Path, max_bytes: int = TEXT_READ_MAX_BYTES) -> tuple[str, str]:
    data = path.read_bytes()[:max_bytes]
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "replace"


def _image_preview(path: Path, mime_type: str, size_bytes: int) -> tuple[str, str, str]:
    if size_bytes > IMAGE_PREVIEW_MAX_BYTES:
        return (
            "preview_only",
            "",
            "图片超过 8 MB，未生成内嵌预览，也不会自动做视觉识别。",
        )
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return "preview_only", "", f"data:{mime_type};base64,{encoded}"


def _pdf_summary(path: Path) -> tuple[str, str, str]:
    try:
        from pypdf import PdfReader
    except Exception:
        return "unsupported", "", "文件已收到，但暂时不能读取内容：缺少 PDF 解析组件。"
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for index, page in enumerate(reader.pages[:5], start=1):
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(f"第 {index} 页：\n{text}")
        if not parts:
            return "unreadable", "", "文件已收到，但暂时不能读取内容：PDF 可能是扫描图片。"
        return "ready", _clip_summary("\n\n".join(parts)), ""
    except Exception:
        return "unreadable", "", "文件已收到，但暂时不能读取内容：PDF 解析失败。"


def _text_summary(path: Path) -> tuple[str, str, str]:
    try:
        text, encoding = _read_text_prefix(path)
    except OSError:
        return "unreadable", "", "文件已收到，但暂时不能读取文本内容。"
    return "ready", _clip_summary(f"文本编码：{encoding}\n\n{text}"), ""


def _csv_summary(path: Path) -> tuple[str, str, str]:
    try:
        text, encoding = _read_text_prefix(path)
        rows = list(csv.reader(io.StringIO(text)))
    except Exception:
        return "unreadable", "", "文件已收到，但 CSV 内容暂时不能解析。"
    if not rows:
        return "ready", "CSV 文件为空。", ""
    header = rows[0]
    preview_rows = rows[1:21]
    max_cols = max((len(row) for row in rows[:200]), default=len(header))
    lines = [
        f"CSV 编码：{encoding}",
        f"列数估算：{max_cols}",
        f"样本行数：{max(0, len(rows) - 1)}",
        "列名：" + ("、".join(header) if header else "无"),
    ]
    if preview_rows:
        lines.append("前 20 行：")
        for index, row in enumerate(preview_rows, start=1):
            lines.append(f"{index}. " + " | ".join(cell for cell in row[:12]))
    return "ready", _clip_summary("\n".join(lines)), ""


def _docx_summary(path: Path) -> tuple[str, str, str]:
    try:
        from docx import Document
    except Exception:
        return "unsupported", "", "文件已收到，但暂时不能读取内容：缺少 Word 文档解析组件。"
    try:
        document = Document(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
            if len(blocks) >= 200:
                break
        if len(blocks) < 200:
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
                    if len(blocks) >= 200:
                        break
                if len(blocks) >= 200:
                    break
        if not blocks:
            return "unreadable", "", "文件已收到，但 Word 文档里暂时没有可读取文本。"
        return "ready", _clip_summary("\n".join(blocks)), ""
    except Exception:
        return "unreadable", "", "文件已收到，但 Word 文档解析失败。"


def _xlsx_summary(path: Path) -> tuple[str, str, str]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return "unsupported", "", "文件已收到，但暂时不能读取内容：缺少 Excel 表格解析组件。"
    workbook = None
    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        lines: list[str] = ["Excel 公式仅显示缓存值。"]
        for sheet_name in workbook.sheetnames[:5]:
            sheet = workbook[sheet_name]
            lines.append(f"Sheet：{sheet_name}")
            for row_index, row in enumerate(
                sheet.iter_rows(max_row=30, max_col=12, values_only=True),
                start=1,
            ):
                values = ["" if value is None else str(value) for value in row]
                if any(value for value in values):
                    lines.append(f"{row_index}. " + " | ".join(values))
        if len(lines) == 1:
            return "unreadable", "", "文件已收到，但 Excel 表格里暂时没有可读取内容。"
        return "ready", _clip_summary("\n".join(lines)), ""
    except Exception:
        return "unreadable", "", "文件已收到，但 Excel 表格解析失败。"
    finally:
        if workbook is not None:
            workbook.close()


def summarize_attachment_file(path: Path, mime_type: str, size_bytes: int) -> tuple[str, str, str, str]:
    suffix = path.suffix.lower()
    if size_bytes > SUMMARY_TOO_LARGE_BYTES and suffix not in IMAGE_EXTS:
        return "too_large", "", "", "文件已收到，但文件较大，暂时不生成内容摘要。"

    if suffix in IMAGE_EXTS or mime_type.startswith("image/"):
        status, summary, preview_or_reason = _image_preview(path, mime_type, size_bytes)
        if preview_or_reason.startswith("data:"):
            return status, summary, preview_or_reason, IMAGE_PREVIEW_ONLY_REASON
        return status, summary, "", preview_or_reason
    if suffix in PDF_EXTS or mime_type == "application/pdf":
        status, summary, reason = _pdf_summary(path)
        return status, summary, "", reason
    if suffix in CSV_EXTS or mime_type == "text/csv":
        status, summary, reason = _csv_summary(path)
        return status, summary, "", reason
    if suffix in TEXT_EXTS or mime_type.startswith("text/"):
        status, summary, reason = _text_summary(path)
        return status, summary, "", reason
    if suffix in DOCX_EXTS:
        status, summary, reason = _docx_summary(path)
        return status, summary, "", reason
    if suffix in XLSX_EXTS:
        status, summary, reason = _xlsx_summary(path)
        return status, summary, "", reason
    if suffix in ARCHIVE_EXTS:
        return "unsupported", "", "", "压缩包已收到，当前版本不会解压读取内容。"
    if suffix in AUDIO_EXTS or mime_type.startswith("audio/"):
        return "unsupported", "", "", "音频已收到，当前版本暂不做转写。"
    if suffix in VIDEO_EXTS or mime_type.startswith("video/"):
        return "unsupported", "", "", "视频已收到，当前版本暂不做转写。"
    return "unsupported", "", "", "文件已收到，但暂时不能读取内容：文件类型暂不支持。"


def _is_image_attachment(attachment: dict[str, Any]) -> bool:
    mime_type = str(attachment.get("mime_type") or "")
    file_name = str(attachment.get("file_name") or "")
    return mime_type.startswith("image/") or Path(file_name).suffix.lower() in IMAGE_EXTS


async def recognize_image_attachments(
    attachments: list[dict[str, Any]],
    *,
    paths: RuntimePaths | None = None,
) -> list[dict[str, Any]]:
    runtime_paths = paths or ensure_runtime_dirs()
    updated_attachments: list[dict[str, Any]] = []
    for attachment in attachments:
        if not _is_image_attachment(attachment):
            updated_attachments.append(attachment)
            continue
        preview_data_url = str(attachment.get("preview_data_url") or "")
        if not preview_data_url:
            updated_attachments.append(attachment)
            continue

        from .chat_client import describe_image_data_url

        result = await describe_image_data_url(
            preview_data_url,
            file_name=str(attachment.get("file_name") or "图片"),
            paths=runtime_paths,
        )
        if result.get("ok"):
            next_attachment = conversations.update_attachment_summary(
                str(attachment["id"]),
                summary_status="recognized",
                summary_text=str(result.get("summary") or "").strip(),
                preview_data_url=preview_data_url,
                reason_cn="",
                paths=runtime_paths,
            )
        else:
            next_attachment = conversations.update_attachment_summary(
                str(attachment["id"]),
                summary_status="preview_only",
                summary_text="",
                preview_data_url=preview_data_url,
                reason_cn=str(result.get("message") or IMAGE_PREVIEW_ONLY_REASON),
                paths=runtime_paths,
            )
        updated_attachments.append(next_attachment or attachment)
    return updated_attachments


def register_message_attachments(
    *,
    message_id: str,
    conversation_id: str,
    media_urls: list[str],
    media_types: list[str] | None = None,
    source: str = "weixin",
    paths: RuntimePaths | None = None,
) -> list[dict[str, Any]]:
    runtime_paths = paths or ensure_runtime_dirs()
    attachments: list[dict[str, Any]] = []
    media_types = media_types or []
    for index, raw_path in enumerate(media_urls):
        source_path = assert_safe_attachment_path(raw_path, runtime_paths)
        mime_type = _guess_mime(source_path, media_types[index] if index < len(media_types) else "")
        size_bytes = source_path.stat().st_size
        attachment_id = f"att_{secrets.token_hex(8)}"
        stored_path = _copy_to_attachment_dir(source_path, attachment_id, runtime_paths)
        attachment = conversations.create_attachment_record(
            attachment_id=attachment_id,
            message_id=message_id,
            conversation_id=conversation_id,
            safe_path=stored_path,
            file_name=_safe_name(source_path.name),
            mime_type=mime_type,
            size_bytes=size_bytes,
            summary_status="pending",
            metadata={"source": source, "original_ext": source_path.suffix.lower()},
            paths=runtime_paths,
        )
        status, summary, preview, reason = summarize_attachment_file(stored_path, mime_type, size_bytes)
        updated = conversations.update_attachment_summary(
            attachment["id"],
            summary_status=status,
            summary_text=summary,
            preview_data_url=preview,
            reason_cn=reason,
            paths=runtime_paths,
        )
        attachments.append(updated or attachment)
    return attachments


def register_generated_attachment(
    source_path: str | Path,
    *,
    message_text: str = "已生成文件。",
    conversation_id: str = conversations.PERSONAL_CONVERSATION_ID,
    paths: RuntimePaths | None = None,
) -> dict[str, Any]:
    runtime_paths = paths or ensure_runtime_dirs()
    source = assert_safe_attachment_path(source_path, runtime_paths)
    message = conversations.create_message(
        conversation_id=conversation_id,
        source="assistant",
        role="assistant",
        text=message_text,
        status="sent",
        metadata={"kind": "generated_attachment"},
        emit_event=False,
        paths=runtime_paths,
    )
    attachment = register_message_attachments(
        message_id=message["id"],
        conversation_id=conversation_id,
        media_urls=[str(source)],
        media_types=[],
        source="generated",
        paths=runtime_paths,
    )[0]
    conversations.record_message_event(message["id"], paths=runtime_paths)
    return attachment


def attachment_summaries_for_prompt(attachments: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    for attachment in attachments:
        name = str(attachment.get("file_name") or "附件")
        status = str(attachment.get("summary_status") or "")
        summary = str(attachment.get("summary_text") or "").strip()
        reason = str(attachment.get("reason_cn") or "").strip()
        if summary:
            if status == "recognized":
                lines.append(f"图片 {name} 视觉识别结果：\n{summary}")
            else:
                lines.append(f"附件 {name} 摘要：\n{summary}")
        elif reason:
            lines.append(f"附件 {name}：{reason}")
        else:
            lines.append(f"附件 {name} 已收到，摘要状态：{status or '未知'}。")
    return "\n\n".join(lines)
